from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, Http404, JsonResponse, HttpResponseRedirect, FileResponse
from django.views.generic import ListView, DetailView, TemplateView, CreateView, DeleteView, UpdateView
from django.db.models import Q
from datetime import datetime, date
from django.core.exceptions import ObjectDoesNotExist
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.urls import reverse, reverse_lazy
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import markdown
import os

from .models import Student, Supervisor, DiplomaProject, Group, GroupOrder
from .models import OrderTemplate, TemplateSection, GeneratedDocument, DocumentCollaborator, DocumentHistory
from .forms import StudentSearchForm, OrderGenerationForm, GroupOrderForm
from .forms import OrderTemplateForm, TemplateSectionForm, DocumentGeneratorForm, DocumentCollaboratorForm, DocumentEditForm

class HomeView(TemplateView):
    """Главная страница"""
    template_name = 'diploma_orders/home.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['total_students'] = Student.objects.count()
        context['total_supervisors'] = Supervisor.objects.count()
        context['total_groups'] = Group.objects.count()
        context['recent_students'] = Student.objects.all().order_by('-id')[:5]
        context['groups'] = Group.objects.all().order_by('course', 'name')
        return context

class StudentListView(ListView):
    """Список студентов с фильтрацией"""
    model = Student
    template_name = 'diploma_orders/student_list.html'
    context_object_name = 'students'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = Student.objects.all()
        
        # Поиск по ФИО или номеру студбилета
        query = self.request.GET.get('query')
        if query:
            queryset = queryset.filter(
                Q(last_name__icontains=query) |
                Q(first_name__icontains=query) |
                Q(patronymic__icontains=query) |
                Q(student_id__icontains=query) |
                Q(email__icontains=query)
            )
        
        # Фильтрация по руководителю
        supervisor_id = self.request.GET.get('supervisor')
        if supervisor_id:
            queryset = queryset.filter(
                diploma_project__supervisor_id=supervisor_id
            )
        
        # Фильтрация по группе
        group_id = self.request.GET.get('group')
        if group_id:
            queryset = queryset.filter(group_id=group_id)
        
        # Фильтрация по статусу
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(diploma_project__status=status)
        
        return queryset.select_related(
            'diploma_project', 
            'diploma_project__supervisor',
            'group'
        ).only(  # Добавляем только нужные поля
            'id', 'last_name', 'first_name', 'patronymic', 
            'student_id', 'email', 'phone', 'photo', 'group_id',
            'diploma_project__topic', 
            'diploma_project__status',
            'diploma_project__deadline',
            'diploma_project__supervisor__last_name',
            'diploma_project__supervisor__first_name',
            'diploma_project__supervisor__patronymic'
        )
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Создаем форму с текущими параметрами фильтрации
        initial_data = {
            'query': self.request.GET.get('query', ''),
            'supervisor': self.request.GET.get('supervisor', ''),
            'group': self.request.GET.get('group', ''),
            'status': self.request.GET.get('status', ''),
        }
        
        context['search_form'] = StudentSearchForm(initial=initial_data)
        context['supervisors'] = Supervisor.objects.all()
        context['groups'] = Group.objects.all()
        return context

class StudentDetailView(DetailView):
    """Детальная страница студента"""
    model = Student
    template_name = 'diploma_orders/student_detail.html'
    context_object_name = 'student'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['order_form'] = OrderGenerationForm()
        context['today'] = datetime.now().date()
        return context

class GroupListView(ListView):
    """Список групп"""
    model = Group
    template_name = 'diploma_orders/group_list.html'
    context_object_name = 'groups'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Добавляем количество студентов в каждой группе
        groups = context['groups']
        for group in groups:
            students = group.students.all()
            group.student_count = students.count()
            group.with_diploma = students.filter(diploma_project__isnull=False).count()
            group.without_diploma = students.filter(diploma_project__isnull=True).count()
            
            # Вычисляем процент
            if group.student_count > 0:
                group.diploma_percentage = int((group.with_diploma / group.student_count) * 100)
            else:
                group.diploma_percentage = 0
                
        return context

class GroupDetailView(DetailView):
    """Детальная страница группы"""
    model = Group
    template_name = 'diploma_orders/group_detail.html'
    context_object_name = 'group'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        students = self.object.students.all().select_related('diploma_project', 'diploma_project__supervisor')
        context['students'] = students
        context['with_diploma'] = students.filter(diploma_project__isnull=False).count()
        context['without_diploma'] = students.filter(diploma_project__isnull=True).count()
        
        # Вычисляем процент для прогресс бара
        student_count = students.count()
        if student_count > 0:
            context['diploma_percentage'] = int((context['with_diploma'] / student_count) * 100)
        else:
            context['diploma_percentage'] = 0
            
        return context

def generate_order(request, student_id):
    """Генерация приказа в разных форматах"""
    student = get_object_or_404(Student, id=student_id)
    
    try:
        diploma_project = student.diploma_project
        supervisor = diploma_project.supervisor
    except ObjectDoesNotExist:
        messages.error(request, "У студента нет дипломного проекта или руководителя")
        return redirect('diploma_orders:student_detail', pk=student_id)
    
    # Получаем данные для приказа
    order_data = {
        'student_name': student.get_full_name(),
        'student_id': student.student_id,
        'topic': diploma_project.topic,
        'supervisor_name': supervisor.get_full_name(),
        'supervisor_degree': supervisor.academic_degree,
        'supervisor_position': supervisor.position,
        'registration_date': diploma_project.registration_date,
        'deadline': diploma_project.deadline,
        'order_number': f"ДП-{student_id}-{datetime.now().strftime('%Y%m%d')}",
        'current_date': datetime.now().date(),
    }
    
    # Форматируем даты для отображения
    from django.utils.formats import date_format
    registration_date_str = date_format(order_data['registration_date'], format='d E Y г.', use_l10n=True)
    deadline_str = date_format(order_data['deadline'], format='d E Y г.', use_l10n=True)
    current_date_str = date_format(order_data['current_date'], format='d E Y г.', use_l10n=True)
    
    # Шаблон приказа
    order_text = f"""ПРИКАЗ № {order_data['order_number']}
О закреплении темы дипломной работы

На основании решения кафедры и в соответствии с учебным планом,

ПРИКАЗЫВАЮ:

1. Закрепить за студентом {order_data['student_name']} (студ. билет № {order_data['student_id']}) тему дипломной работы:
   "{order_data['topic']}"

2. Назначить научным руководителем {order_data['supervisor_name']}, {order_data['supervisor_degree']}, {order_data['supervisor_position']}.

3. Установить срок сдачи дипломной работы: {deadline_str}.

Тема утверждена: {registration_date_str}.

Декан факультета:
___________________ /                    /

{current_date_str} г.
"""
    
    # Если это POST запрос
    if request.method == 'POST':
        format_choice = request.POST.get('format', 'preview')
        
        if format_choice == 'preview':
            # Отображение в браузере
            return render(request, 'diploma_orders/order_preview.html', {
                'student': student,
                'order_text': order_text,
                'order_data': order_data,
                'registration_date_str': registration_date_str,
                'deadline_str': deadline_str,
                'current_date_str': current_date_str,
            })
        
        elif format_choice == 'md':
            # Генерация Markdown
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            md_content = f"""# ПРИКАЗ № {order_data['order_number']}
## О закреплении темы дипломной работы

**Студент:** {order_data['student_name']} (№ {order_data['student_id']})

**Тема дипломной работы:** "{order_data['topic']}"

**Научный руководитель:** {order_data['supervisor_name']}, {order_data['supervisor_degree']}, {order_data['supervisor_position']}

**Дата регистрации темы:** {registration_date_str}

**Срок сдачи:** {deadline_str}

---
*Текст приказа:*

{order_text}
"""
            
            response = HttpResponse(md_content, content_type='text/markdown')
            response['Content-Disposition'] = f'attachment; filename="приказ_{student.student_id}.md"'
            return response
        
        elif format_choice == 'docx':
            # Генерация Word документа
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            document = Document()
            
            # Заголовок
            title = document.add_heading(f'ПРИКАЗ № {order_data["order_number"]}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            document.add_heading('О закреплении темы дипломной работы', 1)
            document.add_paragraph()
            
            # Информация о студенте
            p = document.add_paragraph()
            p.add_run('Студент: ').bold = True
            p.add_run(f'{order_data["student_name"]} (студ. билет № {order_data["student_id"]})')
            
            p = document.add_paragraph()
            p.add_run('Тема дипломной работы: ').bold = True
            p.add_run(f'"{order_data["topic"]}"')
            
            p = document.add_paragraph()
            p.add_run('Научный руководитель: ').bold = True
            p.add_run(f'{order_data["supervisor_name"]}, {order_data["supervisor_degree"]}, {order_data["supervisor_position"]}')
            
            p = document.add_paragraph()
            p.add_run('Дата регистрации темы: ').bold = True
            p.add_run(registration_date_str)
            
            p = document.add_paragraph()
            p.add_run('Срок сдачи: ').bold = True
            p.add_run(deadline_str)
            
            document.add_paragraph()
            document.add_heading('Текст приказа', 2)
            
            # Добавление полного текста приказа
            for line in order_text.split('\n'):
                if line.startswith('ПРИКАЗЫВАЮ:'):
                    p = document.add_paragraph(line)
                    p.runs[0].bold = True
                else:
                    document.add_paragraph(line)
            
            # Сохранение в поток
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            
            response = HttpResponse(
                file_stream.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="приказ_{student.student_id}.docx"'
            return response
    
    # По умолчанию показываем форму (для GET запроса)
    return render(request, 'diploma_orders/student_detail.html', {
        'student': student,
        'order_form': OrderGenerationForm(),
        'today': datetime.now().date(),
    })
    from .models import GroupOrder

class GroupOrderListView(ListView):
    """Список приказов по группам"""
    model = GroupOrder
    template_name = 'diploma_orders/group_order_list.html'
    context_object_name = 'orders'
    paginate_by = 20
    
    def get_queryset(self):
        return GroupOrder.objects.all().select_related('group').order_by('-order_date')

def create_group_order(request, group_id):
    """Создание приказа по группе"""
    group = get_object_or_404(Group, id=group_id)
    
    if request.method == 'POST':
        form = GroupOrderForm(request.POST)
        if form.is_valid():
            # Генерируем номер приказа
            from datetime import datetime
            order_number = f"УП-{datetime.now().strftime('%Y%m%d')}-{datetime.now().strftime('%H%M%S')}"
            
            # Создаем приказ
            order = GroupOrder.objects.create(
                group=group,
                order_number=order_number,
                order_date=form.cleaned_data['order_date'],
                study_form=form.cleaned_data['study_form'],
                direction=form.cleaned_data['direction'],
                note=form.cleaned_data.get('note', '')
            )
            
            messages.success(request, f'Приказ №{order_number} успешно создан!')
            return redirect('diploma_orders:group_order_detail', order_id=order.id)
    else:
        form = GroupOrderForm()
    
    return render(request, 'diploma_orders/create_group_order.html', {
        'group': group,
        'form': form,
    })

def group_order_detail(request, order_id):
    """Детальная страница приказа по группе"""
    order = get_object_or_404(GroupOrder, id=order_id)
    students = order.group.students.all().select_related('diploma_project', 'diploma_project__supervisor')
    
    return render(request, 'diploma_orders/group_order_detail.html', {
        'order': order,
        'group': order.group,
        'students': students,
    })

def generate_group_order_docx(request, order_id):
    """Генерация приказа по группе в формате DOCX"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime
    
    order = get_object_or_404(GroupOrder, id=order_id)
    students = order.group.students.all().select_related('diploma_project', 'diploma_project__supervisor')
    
    # Создаем документ
    document = Document()
    
    # Настройка страницы
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    
    # Шрифт Times New Roman
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # 1. Заголовок - Министерство науки и высшего образования РФ
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Министерство науки и высшего образования Российской Федерации')
    run.bold = True
    run.font.size = Pt(14)
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 2. Учреждение
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Федеральное государственное бюджетное образовательное учреждение\nвысшего образования')
    run.bold = True
    run.font.size = Pt(14)
    
    # Интервал 0.5 см
    document.add_paragraph().paragraph_format.space_after = Pt(14)
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Государственный университет управления»')
    run.bold = True
    run.font.size = Pt(16)
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 3. ПРИКАЗ
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ПРИКАЗ')
    run.bold = True
    run.font.size = Pt(16)
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 4. Дата, город и номер приказа
    from django.utils.formats import date_format
    order_date_str = date_format(order.order_date, format='«d» E Y г.')
    
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Настройка ширины колонок
    widths = [Inches(2), Inches(3), Inches(2)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    row = table.rows[0]
    row.cells[0].text = order_date_str
    row.cells[1].text = 'г. Москва'
    row.cells[2].text = f'№ {order.order_number}'
    
    # Выравнивание
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 5. Заголовок приказа
    p = document.add_paragraph()
    run = p.add_run('ОБ утверждении тем выпускных\nквалификационных работ и назначении руководителей')
    run.bold = True
    run.font.size = Pt(14)
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 6. ПРИКАЗЫВАЮ:
    p = document.add_paragraph()
    run = p.add_run('ПРИКАЗЫВАЮ:')
    run.bold = True
    run.font.size = Pt(14)
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 7. Текст приказа
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'1. Утвердить темы выпускных квалификационных работ и назначить научных руководителей для студентов группы {order.group.name} ({order.get_study_form_display()} форма обучения) направления подготовки {order.direction}:')
    run.font.size = Pt(14)
    
    # 8. Таблица с данными
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = 'ФИО студента, тема ВКР'
    hdr_cells[2].text = 'Научный руководитель'
    
    # Настройка ширины колонок
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(10)
    table.columns[2].width = Cm(6)
    
    # Добавляем студентов
    for i, student in enumerate(students, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        
        if student.diploma_project:
            student_info = f'{student.get_full_name()}\nТема: {student.diploma_project.topic}'
            supervisor_info = f'{student.diploma_project.supervisor.get_full_name()},\n{student.diploma_project.supervisor.academic_degree},\n{student.diploma_project.supervisor.position}'
        else:
            student_info = f'{student.get_full_name()}\nТема: не назначена'
            supervisor_info = 'не назначен'
        
        row_cells[1].text = student_info
        row_cells[2].text = supervisor_info
    
    # Интервал 1 см
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    # 9. Пункт 2 приказа
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('2. Контроль за исполнением настоящего приказа возложить на и.о. заведующего кафедрой информационных систем Д.В. Стефановского.')
    run.font.size = Pt(14)
    
    # Интервал 2 см для подписей
    document.add_paragraph().paragraph_format.space_after = Pt(56)
    
    # 10. Подписи
    # Левая колонка
    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Настройка ширины колонок
    table.columns[0].width = Cm(9)
    table.columns[1].width = Cm(9)
    
    # Убираем границы у таблицы
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # Заполняем таблицу
    rows = table.rows
    
    # Строка 1: Проректор слева
    rows[0].cells[0].text = 'Проректор'
    rows[0].cells[1].text = ''
    
    # Строка 2: Подпись слева
    rows[1].cells[0].text = '___________________ Д.Ю. Брюханов'
    rows[1].cells[1].text = ''
    
    # Пустая строка
    rows[2].cells[0].text = ''
    rows[2].cells[1].text = ''
    
    # Строка 3: Проект приказа вносит
    rows[3].cells[0].text = 'Проект приказа вносит:'
    rows[3].cells[1].text = 'Согласовано:'
    
    # Строка 4: Должности
    rows[4].cells[0].text = 'И.о. заведующего кафедрой\nинформационных систем'
    rows[4].cells[1].text = 'И.о. директора Института\nинформационных систем'
    
    # Строка 5: Подписи
    rows[5].cells[0].text = '___________________ Д.В. Стефановский'
    rows[5].cells[1].text = '___________________ О.М. Писарева'
    
    # Добавляем еще подписи справа
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    table2 = document.add_table(rows=3, cols=1)
    table2.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table2.columns[0].width = Cm(9)
    
    # Убираем границы
    tbl2 = table2._tbl
    tblPr2 = tbl2.tblPr
    tblBorders2 = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders2.append(border)
    
    tblPr2.append(tblBorders2)
    
    rows2 = table2.rows
    rows2[0].cells[0].text = 'Заместитель директора\nПравового департамента'
    rows2[1].cells[0].text = '___________________ В.В. Андросенко'
    rows2[2].cells[0].text = ''
    
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    
    table3 = document.add_table(rows=3, cols=1)
    table3.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table3.columns[0].width = Cm(9)
    
    # Убираем границы
    tbl3 = table3._tbl
    tblPr3 = tbl3.tblPr
    tblBorders3 = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders3.append(border)
    
    tblPr3.append(tblBorders3)
    
    rows3 = table3.rows
    rows3[0].cells[0].text = 'Директор Департамента академической политики\nи реализации образовательных программ'
    rows3[1].cells[0].text = '___________________ Н.А. Стракова'
    rows3[2].cells[0].text = ''
    
    # Сохраняем документ
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    # Отправляем файл
    response = HttpResponse(
        file_stream.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="приказ_группа_{order.group.name}_{order.order_number}.docx"'
    
    return response

def generate_group_order_preview(request, order_id):
    """Предпросмотр приказа по группе"""
    order = get_object_or_404(GroupOrder, id=order_id)
    students = order.group.students.all().select_related('diploma_project', 'diploma_project__supervisor')
    
    # Форматируем дату
    from django.utils.formats import date_format
    order_date_str = date_format(order.order_date, format='«d» E Y г.')
    
    return render(request, 'diploma_orders/group_order_preview.html', {
        'order': order,
        'group': order.group,
        'students': students,
        'order_date_str': order_date_str,
    })
# ... существующий код ...

from .models import OrderTemplate, TemplateSection, GeneratedDocument, DocumentCollaborator
from .forms import OrderTemplateForm, TemplateSectionForm, DocumentGeneratorForm, DocumentCollaboratorForm, DocumentEditForm

# === Умный редактор документов ===

class TemplateListView(ListView):
    """Список шаблонов"""
    model = OrderTemplate
    template_name = 'diploma_orders/template_list.html'
    context_object_name = 'templates'
    
    def get_queryset(self):
        return OrderTemplate.objects.filter(is_active=True)


class TemplateDetailView(DetailView):
    """Детальный просмотр шаблона"""
    model = OrderTemplate
    template_name = 'diploma_orders/template_detail.html'
    context_object_name = 'template'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['sections'] = self.object.sections.all()
        context['section_form'] = TemplateSectionForm()
        return context


def template_editor(request, template_id):
    """Редактор шаблона с перетаскиванием разделов"""
    template = get_object_or_404(OrderTemplate, id=template_id)
    sections = template.sections.all()
    
    if request.method == 'POST':
        if 'save_sections' in request.POST:
            # Сохраняем порядок разделов
            section_order = json.loads(request.POST.get('section_order', '[]'))
            for idx, section_id in enumerate(section_order):
                TemplateSection.objects.filter(id=section_id, template=template).update(order=idx)
            messages.success(request, 'Порядок разделов сохранен!')
            return redirect('diploma_orders:template_editor', template_id=template.id)
        
        elif 'add_section' in request.POST:
            form = TemplateSectionForm(request.POST)
            if form.is_valid():
                section = form.save(commit=False)
                section.template = template
                section.save()
                messages.success(request, f'Раздел "{section.title}" добавлен!')
                return redirect('diploma_orders:template_editor', template_id=template.id)
    
    return render(request, 'diploma_orders/template_editor.html', {
        'template': template,
        'sections': sections,
        'section_form': TemplateSectionForm(),
    })


def generate_document(request, object_type, object_id):
    """Генерация документа с умным редактором"""
    # Определяем объект
    if object_type == 'student':
        obj = get_object_or_404(Student, id=object_id)
        obj_name = obj.get_full_name()
    elif object_type == 'group':
        obj = get_object_or_404(Group, id=object_id)
        obj_name = obj.name
    else:
        raise Http404
    
    # Получаем доступные шаблоны
    templates = OrderTemplate.objects.filter(is_active=True)
    
    if request.method == 'POST':
        template_id = request.POST.get('template')
        template = get_object_or_404(OrderTemplate, id=template_id)
        
        # Собираем данные
        data = {}
        for field in template.get_available_fields_list():
            data[field] = request.POST.get(field, '')
        
        # Добавляем системные данные
        data.update({
            'object_type': object_type,
            'object_id': object_id,
            'object_name': obj_name,
            'current_date': datetime.now().strftime('%d.%m.%Y'),
            'user': request.user.get_full_name() if request.user.is_authenticated else 'Система',
        })
        
        # Генерируем контент
        content = template.content
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            content = content.replace(placeholder, str(value))
        
        # Создаем документ
        doc_number = f"DOC-{object_type.upper()}-{object_id}-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        
        document = GeneratedDocument.objects.create(
            template=template,
            student=obj if object_type == 'student' else None,
            group=obj if object_type == 'group' else None,
            document_data=data,
            content=content,
            document_number=doc_number,
            document_date=datetime.now().date(),
            created_by=request.user if request.user.is_authenticated else None,
            status='draft'
        )
        
        # Добавляем создателя как редактора
        if request.user.is_authenticated:
            DocumentCollaborator.objects.create(
                document=document,
                user=request.user,
                role='editor',
                can_edit=True,
                can_comment=True,
                can_approve=True,
                can_sign=False
            )
        
        messages.success(request, f'Документ {doc_number} создан!')
        return redirect('diploma_orders:document_edit', document_id=document.id)
    
    return render(request, 'diploma_orders/document_generator.html', {
        'object': obj,
        'object_type': object_type,
        'templates': templates,
        'form': DocumentGeneratorForm(),
    })


def document_edit(request, document_id):
    """Редактирование документа в реальном времени"""
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    # Проверяем права
    if request.user.is_authenticated:
        is_collaborator = DocumentCollaborator.objects.filter(
            document=document,
            user=request.user,
            can_edit=True
        ).exists()
        
        if not is_collaborator and document.created_by != request.user:
            messages.error(request, 'У вас нет прав для редактирования этого документа')
            return redirect('diploma_orders:document_view', document_id=document_id)
    
    if request.method == 'POST':
        form = DocumentEditForm(request.POST, instance=document)
        if form.is_valid():
            form.save()
            
            # Сохраняем историю изменений
            DocumentHistory.objects.create(
                document=document,
                user=request.user if request.user.is_authenticated else None,
                action='edit',
                changes=request.POST.get('changes', '')
            )
            
            messages.success(request, 'Изменения сохранены!')
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'status': 'success', 'message': 'Сохранено'})
            return redirect('diploma_orders:document_edit', document_id=document.id)
    
    else:
        form = DocumentEditForm(instance=document)
    
    collaborators = document.collaborators.all().select_related('user')
    collaborator_form = DocumentCollaboratorForm()
    
    return render(request, 'diploma_orders/document_editor.html', {
        'document': document,
        'form': form,
        'collaborators': collaborators,
        'collaborator_form': collaborator_form,
    })


def add_collaborator(request, document_id):
    """Добавление участника документа"""
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    if request.method == 'POST':
        form = DocumentCollaboratorForm(request.POST)
        if form.is_valid():
            collaborator = form.save(commit=False)
            collaborator.document = document
            collaborator.save()
            
            messages.success(request, f'Участник {collaborator.user.get_full_name()} добавлен!')
            return redirect('diploma_orders:document_edit', document_id=document.id)
    
    return redirect('diploma_orders:document_edit', document_id=document.id)


def remove_collaborator(request, document_id, collaborator_id):
    """Удаление участника документа"""
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    if request.user == document.created_by:
        DocumentCollaborator.objects.filter(
            id=collaborator_id,
            document=document
        ).delete()
        messages.success(request, 'Участник удален!')
    
    return redirect('diploma_orders:document_edit', document_id=document.id)


def export_document(request, document_id, format_type):
    """Экспорт документа в разных форматах"""
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    if format_type == 'html':
        # Генерация HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{document.document_number}</title>
            <style>
                body {{ font-family: 'Times New Roman', serif; font-size: 14pt; }}
                .header {{ text-align: center; margin-bottom: 40px; }}
                .content {{ line-height: 1.6; }}
                .signatures {{ margin-top: 100px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Документ № {document.document_number}</h1>
                <p>от {document.document_date.strftime('%d.%m.%Y')}</p>
            </div>
            <div class="content">
                {document.content}
            </div>
        </body>
        </html>
        """
        
        response = HttpResponse(html_content, content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="{document.document_number}.html"'
        return response
    
    elif format_type == 'docx':
        # Генерация DOCX из шаблона
        if document.template and document.template.docx_template:
            # Используем существующий шаблон DOCX
            from docx import Document as DocxDocument
            import io
            
            doc = DocxDocument(document.template.docx_template.path)
            
            # Заменяем плейсхолдеры
            for paragraph in doc.paragraphs:
                for key, value in document.document_data.items():
                    placeholder = f'{{{key}}}'
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Сохраняем в поток
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            response = HttpResponse(
                file_stream.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{document.document_number}.docx"'
            return response
        else:
            # Генерируем простой DOCX
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            import io
            
            doc = DocxDocument()
            
            # Заголовок
            title = doc.add_heading(f'Документ № {document.document_number}', 0)
            
            # Дата
            doc.add_paragraph(f'от {document.document_date.strftime("%d.%m.%Y")}')
            
            # Содержимое
            for line in document.content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Сохраняем
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            response = HttpResponse(
                file_stream.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{document.document_number}.docx"'
            return response
    
    elif format_type == 'pdf':
        # Генерация PDF (требует библиотеки reportlab или weasyprint)
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            import io
            
            buffer = io.BytesIO()
            p = canvas.Canvas(buffer, pagesize=A4)
            width, height = A4
            
            # Заголовок
            p.setFont("Helvetica-Bold", 16)
            p.drawString(50, height - 50, f"Документ № {document.document_number}")
            
            p.setFont("Helvetica", 12)
            p.drawString(50, height - 80, f"от {document.document_date.strftime('%d.%m.%Y')}")
            
            # Содержимое
            p.setFont("Helvetica", 10)
            y = height - 120
            for line in document.content.split('\n'):
                if y < 50:  # Новая страница
                    p.showPage()
                    y = height - 50
                    p.setFont("Helvetica", 10)
                
                p.drawString(50, y, line[:100])  # Ограничиваем длину строки
                y -= 15
            
            p.save()
            buffer.seek(0)
            
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{document.document_number}.pdf"'
            return response
            
        except ImportError:
            messages.error(request, 'Для генерации PDF установите reportlab')
            return redirect('diploma_orders:document_edit', document_id=document_id)
    
    else:
        raise Http404


def document_history(request, document_id):
    """История изменений документа"""
    document = get_object_or_404(GeneratedDocument, id=document_id)
    history = DocumentHistory.objects.filter(document=document).select_related('user').order_by('-timestamp')
    
    return render(request, 'diploma_orders/document_history.html', {
        'document': document,
        'history': history,
    })


class DocumentListView(ListView):
    """Список всех документов"""
    model = GeneratedDocument
    template_name = 'diploma_orders/document_list.html'
    context_object_name = 'documents'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = GeneratedDocument.objects.all()
        
        # Фильтрация по типу
        doc_type = self.request.GET.get('type')
        if doc_type:
            if doc_type == 'student':
                queryset = queryset.filter(student__isnull=False)
            elif doc_type == 'group':
                queryset = queryset.filter(group__isnull=False)
        
        # Фильтрация по статусу
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(status=status)
        
        # Фильтрация по автору
        if self.request.user.is_authenticated:
            if 'my' in self.request.GET:
                queryset = queryset.filter(created_by=self.request.user)
        
        return queryset.select_related('template', 'student', 'group', 'created_by')
# ... существующий код в diploma_orders/views.py ...

# Добавляем в конец файла:

# === API представления ===

def api_section_detail(request, section_id):
    """API для работы с разделом"""
    section = get_object_or_404(TemplateSection, id=section_id)
    
    if request.method == 'GET':
        # Возвращаем данные раздела в формате JSON
        return JsonResponse({
            'id': section.id,
            'title': section.title,
            'content': section.content,
            'order': section.order,
            'is_required': section.is_required,
            'can_be_deleted': section.can_be_deleted,
            'can_be_edited': section.can_be_edited,
        })
    
    elif request.method == 'POST':
        # Обновление раздела
        if request.user.is_authenticated:
            data = json.loads(request.body)
            section.title = data.get('title', section.title)
            section.content = data.get('content', section.content)
            section.order = data.get('order', section.order)
            section.save()
            
            return JsonResponse({'success': True})
        
        return JsonResponse({'success': False, 'error': 'Не авторизован'})
    
    elif request.method == 'DELETE':
        # Удаление раздела
        if request.user.is_authenticated and section.can_be_deleted:
            section.delete()
            return JsonResponse({'success': True})
        
        return JsonResponse({'success': False, 'error': 'Нельзя удалить этот раздел'})


def api_section_edit_form(request, section_id):
    """Форма редактирования раздела (HTML)"""
    section = get_object_or_404(TemplateSection, id=section_id)
    form = TemplateSectionForm(instance=section)
    
    return render(request, 'diploma_orders/partials/section_edit_form.html', {
        'form': form,
        'section': section,
    })


def save_template_content(request, template_id):
    """Сохранение содержимого шаблона через AJAX"""
    if request.method == 'POST' and request.user.is_authenticated:
        try:
            template = OrderTemplate.objects.get(id=template_id)
            data = json.loads(request.body)
            
            if 'content' in data:
                template.content = data['content']
                template.save()
                return JsonResponse({'success': True})
            
            return JsonResponse({'success': False, 'error': 'Нет данных'})
            
        except OrderTemplate.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Шаблон не найден'})
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Неверный формат данных'})
    
    return JsonResponse({'success': False, 'error': 'Не авторизован'})


def document_view(request, document_id):
    """Просмотр документа"""
    document = get_object_or_404(GeneratedDocument, id=document_id)
    
    return render(request, 'diploma_orders/document_view.html', {
        'document': document,
    })


class TemplateCreateView(CreateView):
    """Создание нового шаблона"""
    model = OrderTemplate
    form_class = OrderTemplateForm
    template_name = 'diploma_orders/template_create.html'
    success_url = reverse_lazy('diploma_orders:template_list')
    
    def form_valid(self, form):
        form.instance.created_by = self.request.user
        return super().form_valid(form)


class SectionDeleteView(DeleteView):
    """Удаление раздела"""
    model = TemplateSection
    template_name = 'diploma_orders/section_confirm_delete.html'
    
    def get_success_url(self):
        return reverse_lazy('diploma_orders:template_editor', 
                          kwargs={'template_id': self.object.template.id})
    
    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        if self.object.can_be_deleted:
            success_url = self.get_success_url()
            self.object.delete()
            return HttpResponseRedirect(success_url)
        else:
            messages.error(request, 'Этот раздел нельзя удалить')
            return redirect('diploma_orders:template_editor', 
                          template_id=self.object.template.id)
# Добавьте эти функции в views.py

def api_template_fields(request, template_id):
    """API для получения полей шаблона"""
    template = get_object_or_404(OrderTemplate, id=template_id)
    
    fields = []
    available_fields = template.get_available_fields_list()
    
    for field_name in available_fields:
        # Автозаполнение значений для студента/группы
        value = ""
        placeholder = ""
        help_text = ""
        
        # Преобразуем имя поля в читаемый формат
        label = field_name.replace('_', ' ').title()
        
        fields.append({
            'name': field_name,
            'label': label,
            'value': value,
            'placeholder': placeholder,
            'help_text': help_text
        })
    
    return JsonResponse({'fields': fields})

def api_template_preview(request, template_id):
    """API для предпросмотра шаблона"""
    if request.method == 'POST':
        template = get_object_or_404(OrderTemplate, id=template_id)
        
        # Собираем данные из формы
        data = {}
        for field in template.get_available_fields_list():
            data[field] = request.POST.get(field, '')
        
        # Добавляем системные данные
        data['current_date'] = datetime.now().strftime('%d.%m.%Y')
        data['generated_date'] = datetime.now().strftime('%d %B %Y г.')
        
        # Генерируем предпросмотр
        content = template.content
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            content = content.replace(placeholder, str(value))
        
        # Форматируем для отображения
        preview_html = f"""
        <div style="font-family: 'Times New Roman', serif; font-size: 14pt; line-height: 1.5;">
            {content.replace('\n', '<br>')}
        </div>
        """
        
        return JsonResponse({
            'success': True,
            'preview': preview_html
        })
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})